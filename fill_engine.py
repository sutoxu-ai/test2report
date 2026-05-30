"""
轻型动力触探检测报告自动生成工具 - 填充引擎 v2.0
基于新模板（2026-05-26），红色变量内容匹配 + 表格坐标双保险
"""
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, re, subprocess, time
from lxml import etree as _ET
_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _all_paragraphs(doc):
    for p in doc.paragraphs:
        yield ('body', p)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    yield (f'table_{ti}_{ri}_{ci}', p)
    for si, section in enumerate(doc.sections):
        header = section.header
        if header:
            for p in header.paragraphs:
                yield (f'header_{si}', p)


def _red_runs(p):
    result = []
    for i, run in enumerate(p.runs):
        try:
            if run.font.color and run.font.color.rgb:
                if str(run.font.color.rgb).upper() == 'FF0000':
                    result.append((i, run, run.text))
        except OSError:
            pass
    return result


def _set_runs_text(p, new_text, red_only=True):
    BLACK = RGBColor(0, 0, 0)
    reds = _red_runs(p)
    if red_only:
        if reds:
            for i, (ri, run, _) in enumerate(reds):
                _safe_set_run_text(run, str(new_text) if i == 0 else '')
                _safe_set_font_color(run, BLACK)
            return True
    else:
        for i, run in enumerate(p.runs):
            _safe_set_run_text(run, str(new_text) if i == 0 else '')
            _safe_set_font_color(run, BLACK)
        return True
    return False


def _safe_set_run_text(run, text):
    try:
        run.text = str(text)
    except OSError:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        ts = run._element.findall('w:t', ns)
        if not ts:
            ts = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        if ts:
            ts[0].text = str(text)
            for t in ts[1:]:
                t.text = ''


def _safe_set_font_color(run, color):
    try:
        run.font.color.rgb = color
    except OSError:
        pass


def _safe_set_paragraph_text(p, text):
    W = _W
    for r_elem in p._element.findall(f'{W}r'):
        p._element.remove(r_elem)
    r_elem = _ET.SubElement(p._element, f'{W}r')
    t_elem = _ET.SubElement(r_elem, f'{W}t')
    t_elem.text = str(text)


def _safe_set_cell_text(cell, text):
    if len(cell.paragraphs) == 0:
        return
    _safe_set_paragraph_text(cell.paragraphs[0], text)


def _ensure_paragraph_runs(p):
    """确保段落至少有一个 run，如果没 run 则创建一个"""
    if len(p.runs) == 0:
        r = _ET.SubElement(p._element, qn('w:r'))
        _ET.SubElement(r, qn('w:rPr'))
        _ET.SubElement(r, qn('w:t')).text = ''


def _merge_table9_soil_cells(sum_table, start_idx, end_idx):
    """合并表9第一列（土层）中连续相同土层的单元格。
    start_idx/end_idx 是数据行索引（0-based，从 summary_data 第0行算起）。
    对应 table row idx = start_idx+1 / end_idx+1（跳过表头行）。
    全程使用 XML 级别操作，避免 python-docx row.cells 缓存/共享问题。
    """
    from lxml import etree as etree2
    tr_elements = sum_table._tbl.findall(qn('w:tr'))

    # 第一行：restart，保留文字
    tr_first = tr_elements[start_idx + 1]
    tc_first = tr_first.findall(qn('w:tc'))[0]
    tcPr_first = tc_first.find(qn('w:tcPr'))
    if tcPr_first is None:
        tcPr_first = etree2.SubElement(tc_first, qn('w:tcPr'))
    vm_exist = tcPr_first.find(qn('w:vMerge'))
    if vm_exist is not None:
        tcPr_first.remove(vm_exist)
    etree2.SubElement(tcPr_first, qn('w:vMerge')).set(qn('w:val'), 'restart')

    # 后续行：continue，清空文字
    for ri in range(start_idx + 1, end_idx + 1):
        tr = tr_elements[ri + 1]
        tc = tr.findall(qn('w:tc'))[0]
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = etree2.SubElement(tc, qn('w:tcPr'))
        vm_exist = tcPr.find(qn('w:vMerge'))
        if vm_exist is not None:
            tcPr.remove(vm_exist)
        etree2.SubElement(tcPr, qn('w:vMerge'))  # 无val=continue
        # 清空文字：操作 w:tc 下的 w:p/w:r/w:t
        for p_elem in tc.findall(qn('w:p')):
            for r_elem in p_elem.findall(qn('w:r')):
                t_elem = r_elem.find(qn('w:t'))
                if t_elem is not None:
                    t_elem.text = ''


# ===== 主填充函数 =====

def fill_document(template_path, output_path, data):
    import re  # 确保函数内可访问
    doc = Document(template_path)
    from lxml import etree as etree2
    BLACK = RGBColor(0, 0, 0)

    # ===== 0.1 空白行清理 =====
    # 注意：封面区域（#@ 标记到"声明："段落之前）依赖空行撑开分页，不能压缩。
    # 只压缩"声明："段落之后的正文区域（即第2页之后）的连续空行。
    cover_end_pi = None
    for pi, p in enumerate(doc.paragraphs):
        if p.text.strip() == '声明：':
            cover_end_pi = pi
            break
    if cover_end_pi is None:
        cover_end_pi = 30  # 保底：跳过前30个段落

    empty_indices = []
    for pi, p in enumerate(doc.paragraphs):
        if pi <= cover_end_pi:
            continue  # 封面区域不压缩
        if not p.text.strip():
            empty_indices.append(pi)
    to_remove = set()
    for i in range(1, len(empty_indices)):
        if empty_indices[i] == empty_indices[i-1] + 1:
            to_remove.add(empty_indices[i])
    for pi in sorted(to_remove, reverse=True):
        p_elem = doc.paragraphs[pi]._element
        p_elem.getparent().remove(p_elem)

    doc.save(output_path)
    doc = Document(output_path)
    all_p = list(_all_paragraphs(doc))

    # ===== 1. 封面第1页 — 段落级别填充 =====
    # 1a. 工程地点
    for loc, p in all_p:
        full = p.text.strip()
        if '工程地点' in full and '宜昌市' in full:
            _set_runs_text(p, data.get('project_location', ''), True)
    # 备用：匹配纯地点文本
    for loc, p in all_p:
        full = p.text.strip()
        if full == '宜昌市伍家岗区橘乡大道' and not loc.startswith('table'):
            _set_runs_text(p, data.get('project_location', ''), True)

    # 1b. 委托单位
    for loc, p in all_p:
        full = p.text.strip()
        if '委托单位' in full and '宜昌市' in full:
            _set_runs_text(p, data.get('client_name', ''), True)
    for loc, p in all_p:
        full = p.text.strip()
        if full == '宜昌市城市建设投资开发有限公司' and not loc.startswith('table'):
            _set_runs_text(p, data.get('client_name', ''), True)

    # 1c. 报告编号 — 封面
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('报告编号：') and 'DT' in full:
            _set_runs_text(p, data.get('report_number', ''), True)
    # 页眉 DT 编号
    for loc, p in all_p:
        if loc.startswith('header') and 'DT' in p.text:
            reds = _red_runs(p)
            for ri, run, _ in reds:
                if 'DT' in run.text:
                    _safe_set_run_text(run, data.get('report_number', ''))
                    _safe_set_font_color(run, BLACK)
                    break

    # 1d. 检测日期 — 封面
    date_str = data.get('test_date', '')
    # 用于正则解析：把顿号替换成横杠，例：'2026、05、08' → '2026-05-08'
    normalized = re.sub(r'[、，]', '-', date_str.strip())

    for loc, p in all_p:
        full = p.text.strip()
        full_nocolon = full.replace('：', ':').replace('\uff1a', ':')
        if full_nocolon.startswith('检测日期:'):
            reds = _red_runs(p)
            if len(reds) >= 6:
                # 用 normalized（顿号已替换为-）解析年月日
                parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', normalized)
                if parts:
                    y, m, d = parts.groups()
                    # reds 对应模板：['20','26','年','05','月','08日'] 共6个
                    texts = ['20', y[2:], '年', m.zfill(2), '月', d.zfill(2) + '日']
                    for i, (ri, run, _) in enumerate(reds):
                        _safe_set_run_text(run, texts[i] if i < len(texts) else '')
                        _safe_set_font_color(run, BLACK)
                else:
                    # 无法解析时，完整日期填入第一个红色 run
                    for i, (ri, run, _) in enumerate(reds):
                        _safe_set_run_text(run, date_str if i == 0 else '')
                        _safe_set_font_color(run, BLACK)
            break

    # 1e. 报告日期 — 封面 + 首页
    # 模板中有多处 YYYY年MM月DD日 红色字段（封面、首页等），需全部填充
    # 用元素 id 做去重，避免基于文本的去重误杀不同位置的同文段落
    report_date = data.get('report_date', '')
    date_handled = set()
    for loc, p in all_p:
        full = p.text.strip()
        elem_id = id(p._element)
        if elem_id in date_handled:
            continue
        if full.startswith('检测日期：') or re.match(r'^\d{4}\.\d{2}\.\d{2}$', full):
            continue
        if re.match(r'^\d{4}年\d{2}月\d{2}日$', full):
            reds = _red_runs(p)
            if reds:
                date_handled.add(elem_id)
                parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', report_date)
                if parts and len(reds) >= 4:
                    y, m, d = parts.groups()
                    texts = [y, '年', m.zfill(2), '月', d.zfill(2), '日']
                    for i, (ri, run, _) in enumerate(reds):
                        _safe_set_run_text(run, texts[i] if i < len(texts) else '')
                        _safe_set_font_color(run, BLACK)
                else:
                    for i, (ri, run, _) in enumerate(reds):
                        _safe_set_run_text(run, report_date if i == 0 else '')
                        _safe_set_font_color(run, BLACK)

    # 1f. 工程名称 — 封面底部 + 第2页顶部（多处在模板中是红色）
    # 匹配"宜昌市共同南路..."开头的段落
    project_name = data.get('project_name', '')
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('宜昌市共同南路'):
            _set_runs_text(p, project_name, True)

    # ===== 2. 短格式日期 "2026.05.08" =====
    for loc, p in all_p:
        full = p.text.strip()
        if re.match(r'^\d{4}\.\d{2}\.\d{2}$', full):
            parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', full)
            if parts:
                y, m, d = parts.groups()
                short = f'{y}.{m.zfill(2)}.{d.zfill(2)}'
            else:
                short = full
            _set_runs_text(p, short, True)

    # ===== 3. 检测单位信息 =====
    test_unit_info = data.get('test_unit_info', '')
    if test_unit_info:
        unit_lines = test_unit_info.strip().split('\n')
        unit_idx = 0
        for loc, p in all_p:
            full = p.text.strip()
            if not loc.startswith('table') and (
                full.startswith('检测单位：') or full.startswith('地    址：') or
                full.startswith('邮    编：') or full.startswith('电    话：') or
                full.startswith('传    真：') or full.startswith('监督电话：')
            ):
                _ensure_paragraph_runs(p)
                if unit_idx < len(unit_lines):
                    for ri, run in enumerate(p.runs):
                        _safe_set_run_text(run, unit_lines[unit_idx] if ri == 0 else '')
                        try:
                            run.font.color.rgb = BLACK
                        except:
                            pass
                    unit_idx += 1

    # ===== 4. 第3页"首页"表（Table 1）=====
    # 此表包含：工程名称、委托单位、检测性质(固定)、检测目的(固定)、检测项目(固定)、
    # 检测依据、地基类型、土层、设计承载力、地基面积、抽检数量、检测日期、
    # 总进尺、检测深度、检测位置(只填前半)、检测结论(只填后半)、备注(固定)
    if len(doc.tables) > 1:
        t_home = doc.tables[1]

        # 4a0. 工程名称 row[0].cell[1-3] — 之前缺失，section 1f 只通过文本匹配覆盖
        project_name = data.get('project_name', '')
        if len(t_home.rows) > 0:
            for ci in range(1, min(len(t_home.rows[0].cells), 4)):
                for p in t_home.rows[0].cells[ci].paragraphs:
                    _set_runs_text(p, project_name, True)

        # 4a1. 委托单位 row[1].cell[1-3] — 之前缺失，section 1b 只覆盖封面 body 段落
        client_name = data.get('client_name', '')
        if len(t_home.rows) > 1:
            for ci in range(1, min(len(t_home.rows[1].cells), 4)):
                for p in t_home.rows[1].cells[ci].paragraphs:
                    _set_runs_text(p, client_name, True)

        # 4a. 检测依据 row[5].cell[1]
        testing_standards_page1 = data.get('testing_standards_page1', '')
        if testing_standards_page1 and len(t_home.rows) > 5 and len(t_home.rows[5].cells) > 1:
            for p in t_home.rows[5].cells[1].paragraphs:
                _set_runs_text(p, testing_standards_page1, True)

        # 4b. 地基类型 row[6].cell[1], 土层 row[6].cell[3]
        if len(t_home.rows) > 6:
            if len(t_home.rows[6].cells) > 1:
                for p in t_home.rows[6].cells[1].paragraphs:
                    _set_runs_text(p, data.get('foundation_type', ''), True)
            if len(t_home.rows[6].cells) > 3:
                for p in t_home.rows[6].cells[3].paragraphs:
                    _set_runs_text(p, data.get('soil_layer', ''), True)

        # 4c. 设计承载力 / 地基面积 row[7].cell[1] + row[7].cell[3]
        caps = data.get('bearing_capacities', '')
        if caps:
            parts = re.split(r'[、,，\s]+', caps)
            formatted = ['≥' + p.strip() if p.strip() and not p.strip().startswith('≥') else p.strip() for p in parts if p.strip()]
            caps_display = '、'.join(formatted)
        else:
            caps_display = ''
        if len(t_home.rows) > 7:
            if len(t_home.rows[7].cells) > 1:
                for p in t_home.rows[7].cells[1].paragraphs:
                    _set_runs_text(p, caps_display, True)
            foundation_area = data.get('foundation_area', '')
            if foundation_area and len(t_home.rows[7].cells) > 3:
                for p in t_home.rows[7].cells[3].paragraphs:
                    _set_runs_text(p, foundation_area, True)

        # 4d. 抽检数量 / 检测日期 row[8].cell[1] + row[8].cell[3]
        if len(t_home.rows) > 8:
            sc = data.get('sample_count', '')
            if sc and len(t_home.rows[8].cells) > 1:
                display_sc = str(sc) + '点' if not str(sc).endswith('点') else str(sc)
                for p in t_home.rows[8].cells[1].paragraphs:
                    _set_runs_text(p, display_sc, True)
            if len(t_home.rows[8].cells) > 3:
                # 用 normalized 解析（顿号已替换为-）
                parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', normalized)
                if parts:
                    y, m, d = parts.groups()
                    short_date = f'{y}.{m.zfill(2)}.{d.zfill(2)}'
                else:
                    short_date = date_str
                for p in t_home.rows[8].cells[3].paragraphs:
                    _set_runs_text(p, short_date, True)

        # 4e. 总进尺 / 检测深度 row[9].cell[1] + row[9].cell[3]
        if len(t_home.rows) > 9:
            if len(t_home.rows[9].cells) > 1:
                for p in t_home.rows[9].cells[1].paragraphs:
                    _set_runs_text(p, data.get('total_depth', ''), True)
            if len(t_home.rows[9].cells) > 3:
                for p in t_home.rows[9].cells[3].paragraphs:
                    _set_runs_text(p, data.get('test_depth_range', ''), True)

        # 4f. 检测位置 — 只填红色部分，保留黑色后缀"（具体点位见附图1）"
        pile_range = data.get('pile_range', '')
        if pile_range and len(t_home.rows) > 10:
            for ci in range(1, min(len(t_home.rows[10].cells), 4)):
                cell = t_home.rows[10].cells[ci]
                for p in cell.paragraphs:
                    reds = _red_runs(p)
                    if reds:
                        for i, (ri, run, _) in enumerate(reds):
                            if i == 0:
                                _safe_set_run_text(run, pile_range)
                            else:
                                _safe_set_run_text(run, '')
                            _safe_set_font_color(run, BLACK)
                    # 无红色 run 的段落不动，保留原有黑色文字

        # 4g. 检测结论 — 只填红色run（后半句），保留黑色前缀
        conclusion = data.get('test_conclusion', '')
        if conclusion and len(t_home.rows) > 11:
            for ci in range(1, min(len(t_home.rows[11].cells), 4)):
                cell = t_home.rows[11].cells[ci]
                for p in cell.paragraphs:
                    reds = _red_runs(p)
                    if reds:
                        for i, (ri, run, _) in enumerate(reds):
                            _safe_set_run_text(run, conclusion if i == 0 else '')
                            _safe_set_font_color(run, BLACK)

        # 4h. 首页备注 — 固定常量，不填充（模板已有黑色文本"检测位置及数量由施工、监理、设计及建设等单位共同确定。"）

    # ===== 5. 第5页"表1 — 项目概况"（Table 2）=====
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        units = data.get('project_units', {})

        # 从旧匹配迁移到行坐标：
        # Row 0: 工程名称 (col 1-3)
        for ci in range(1, min(len(t2.rows[0].cells), 4)):
            for p in t2.rows[0].cells[ci].paragraphs:
                _set_runs_text(p, data.get('project_name', ''), True)

        # Row 1: 工程地点 (col 1-3)
        for ci in range(1, min(len(t2.rows[1].cells), 4)):
            for p in t2.rows[1].cells[ci].paragraphs:
                _set_runs_text(p, data.get('project_location', ''), True)

        # Row 2: 建设单位 (col 1-3)
        for ci in range(1, min(len(t2.rows[2].cells), 4)):
            for p in t2.rows[2].cells[ci].paragraphs:
                _set_runs_text(p, units.get('construction_unit', ''), True)

        # Row 3: 勘察单位 (col 1-3)
        for ci in range(1, min(len(t2.rows[3].cells), 4)):
            for p in t2.rows[3].cells[ci].paragraphs:
                _set_runs_text(p, units.get('survey', ''), True)

        # Row 4: 设计单位 (col 1-3)
        for ci in range(1, min(len(t2.rows[4].cells), 4)):
            for p in t2.rows[4].cells[ci].paragraphs:
                _set_runs_text(p, units.get('design', ''), True)

        # Row 5: 施工单位 (col 1-3)
        for ci in range(1, min(len(t2.rows[5].cells), 4)):
            for p in t2.rows[5].cells[ci].paragraphs:
                _set_runs_text(p, units.get('construction', ''), True)

        # Row 6: 监理单位 (col 1-3)
        for ci in range(1, min(len(t2.rows[6].cells), 4)):
            for p in t2.rows[6].cells[ci].paragraphs:
                _set_runs_text(p, units.get('supervision', ''), True)

        # Row 7: 见证人 (col 1) — 之前错误映射为 construction_unit，现在修复
        if len(t2.rows) > 7 and len(t2.rows[7].cells) > 1:
            for p in t2.rows[7].cells[1].paragraphs:
                _set_runs_text(p, data.get('witness', ''), True)

        # Row 8: 质量监督站 (col 1-3)
        if len(t2.rows) > 8:
            for ci in range(1, min(len(t2.rows[8].cells), 4)):
                for p in t2.rows[8].cells[ci].paragraphs:
                    _set_runs_text(p, units.get('quality_station', ''), True)

        # Row 9: 结构型式 (col 1) / 基础类型 (col 3)
        if len(t2.rows) > 9:
            if len(t2.rows[9].cells) > 1:
                for p in t2.rows[9].cells[1].paragraphs:
                    _set_runs_text(p, data.get('structure_type', ''), True)
            if len(t2.rows[9].cells) > 3:
                for p in t2.rows[9].cells[3].paragraphs:
                    _set_runs_text(p, data.get('base_type', ''), True)

        # Row 10: 设计承载力 (col 1) / 基底高程 (col 3)
        if len(t2.rows) > 10:
            if len(t2.rows[10].cells) > 1:
                for p in t2.rows[10].cells[1].paragraphs:
                    _set_runs_text(p, caps_display, True)
            if len(t2.rows[10].cells) > 3:
                for p in t2.rows[10].cells[3].paragraphs:
                    _set_runs_text(p, data.get('base_elevation', ''), True)

        # Row 11: 检测方法 — 固定为"采用轻型(10kg)动力触探试验"，不动
        # 但模板中此格是红色，需要去掉红色而不是填充新内容
        if len(t2.rows) > 11 and len(t2.rows[11].cells) > 1:
            for p in t2.rows[11].cells[1].paragraphs:
                for run in p.runs:
                    try:
                        if run.font.color and run.font.color.rgb and str(run.font.color.rgb).upper() == 'FF0000':
                            _safe_set_font_color(run, BLACK)
                    except OSError:
                        pass

        # Row 12: 备注（独立字段，支持自由录入）
        # 填充所有值格（col 1-3），避免合并单元格拆分后部分格遗漏
        table1_remark = data.get('table1_remark', '——')
        if len(t2.rows) > 12:
            for ci in range(1, min(len(t2.rows[12].cells), 4)):
                for p in t2.rows[12].cells[ci].paragraphs:
                    _set_runs_text(p, table1_remark, True)

    # ===== 6. 承载力特征值全局匹配 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('≥') and ('100' in full or '200' in full or '120' in full or '130' in full or '150' in full):
            _set_runs_text(p, caps_display, True)

    # ===== 7. 检测桩号范围（正文中）=====
    for loc, p in all_p:
        full = p.text.strip()
        if '（具体点位见附图' in full and not loc.startswith('table'):
            _set_runs_text(p, pile_range, True)
            break

    # ===== 8. 第9页 六、现场检测（抽样数量/深度/总进尺）=====
    for loc, p in all_p:
        full = p.text.strip()
        if '现场检测由委托方进行抽样' in full and not loc.startswith('table'):
            reds = _red_runs(p)
            if len(reds) >= 3:
                _safe_set_run_text(reds[0][1], str(data.get('sample_count', '')))
                _safe_set_run_text(reds[1][1], str(data.get('test_depth_meters', '')))
                _safe_set_run_text(reds[2][1], str(data.get('total_depth', '')))
            break

    # ===== 9. 地质概况 =====
    # 9-pre. 给"二、地质概况"标题段落加 pageBreakBefore，确保它永远从新页第一行开始
    # （模板依赖空行撑页，内容较少时该标题前的空行会溢到新页成为第一行空白）
    for loc, p in all_p:
        if p.text.strip() == '二、地质概况' and not loc.startswith('table'):
            pPr = p._element.find(f'{_W}pPr')
            if pPr is None:
                from docx.oxml import OxmlElement as _OxmlElement2
                pPr = _OxmlElement2('w:pPr')
                p._element.insert(0, pPr)
            pbk = pPr.find(f'{_W}pageBreakBefore')
            if pbk is None:
                from docx.oxml import OxmlElement as _OxmlElement2
                pbk = _OxmlElement2('w:pageBreakBefore')
                pPr.append(pbk)
            pbk.set(f'{_W}val', 'true')
            break

    geo_mode = data.get('geo_mode', 'full')
    geo_layers = data.get('geo_layers', [])
    geo_desc_text = data.get('geo_description', '')

    # 9a. 地质概况段落
    # 匹配 "由中国兵器工业..." 或 "由...提供的《岩土工程勘察报告》" 开头的段落
    for loc, p in all_p:
        full = p.text.strip()
        if ('《岩土工程勘察报告》' in full or (full.startswith('由') and '勘察' in full)) and '本次检测' in full:
            if geo_mode == 'simple':
                # 无地勘：只显示一句话
                spr = data.get('simple_pile_range', '')
                sft = data.get('simple_foundation_type', '')
                ssl = data.get('simple_soil_layer', '')
                simple_text = f'本次检测{spr}，地基类型为{sft}，主要土层为{ssl}。'
                _set_runs_text(p, simple_text, False)  # False = 替换全部run，包括黑色run
            else:
                # 有地勘：使用用户自定义的地质描述
                if geo_desc_text:
                    _apply_superscript(p, geo_desc_text)
                else:
                    # 回退到锚点替换
                    _set_runs_text(p, full, True)
            break

    # 9b. 地质概况表（Table 2）— 对应 docx 中的 Table 3
    # 需要找到表头包含"岩土名称"的那个表格
    geo_table = None
    for ti, table in enumerate(doc.tables):
        if len(table.rows) > 0:
            first_cell_texts = []
            for cell in table.rows[0].cells:
                first_cell_texts.append(cell.text.strip())
            header_str = ''.join(first_cell_texts)
            if '岩土名称' in header_str and '岩土层描述' in header_str:
                geo_table = table
                break

    if geo_table is not None:
        if geo_mode == 'simple':
            # 无地勘：完全删除整个表格（从文档body中移除w:tbl元素）
            tbl_element = geo_table._tbl
            tbl_element.getparent().remove(tbl_element)
        else:
            # 有地勘：填充用户录入的数据
            while len(geo_table.rows) > len(geo_layers) + 1:
                tr = geo_table.rows[-1]._tr
                geo_table._tbl.remove(tr)
            for i, layer in enumerate(geo_layers):
                ri = i + 1
                if ri < len(geo_table.rows):
                    row = geo_table.rows[ri]
                    cell0 = row.cells[0]
                    _set_runs_text(cell0.paragraphs[0], layer.get('name', ''), True)
                    for pi in range(1, len(cell0.paragraphs)):
                        for run in cell0.paragraphs[pi].runs:
                            _safe_set_run_text(run, '')
                    desc = layer.get('description', '')
                    if len(row.cells[1].paragraphs) > 0:
                        p_desc = row.cells[1].paragraphs[0]
                        reds = _red_runs(p_desc)
                        if reds:
                            for j, (rj, run, _) in enumerate(reds):
                                _safe_set_run_text(run, str(desc) if j == 0 else '')
                                _safe_set_font_color(run, BLACK)
                        else:
                            _set_runs_text(p_desc, desc, False)
                else:
                    new_row = geo_table.add_row()
                    _safe_set_cell_text(new_row.cells[0], layer.get('name', ''))
                    _safe_set_cell_text(new_row.cells[1], layer.get('description', ''))

    # 9c. 简化模式：清除"场区岩土层概况"和"表2"段落文本
    if geo_mode == 'simple':
        for loc, p in all_p:
            full = p.text.strip()
            # 清除"场区岩土层概况"段落
            if full == '场区岩土层概况':
                _set_runs_text(p, '', False)
            # 清除"表2"段落（允许前后有空格）
            elif full.replace(' ', '').replace('\u3000', '') == '表2':
                _set_runs_text(p, '', False)

    # ===== 10. 检测依据（第三章）=====
    testing_items = data.get('testing_items', [])
    extra_items = data.get('extra_testing_items', [])
    all_items = testing_items + extra_items
    all_items = [x for x in all_items if x]
    fixed_item_num = data.get('fixed_item_num', 3)

    if all_items:
        item_idx = 0
        last_detection_p = None
        in_section = False
        for loc, p in all_p:
            full = p.text.strip()
            no_space = full.replace(' ', '').replace('\u3000', '')
            if no_space.startswith('三') and '检测依据' in no_space:
                in_section = True
                continue
            if in_section and (no_space and no_space[0] in '四五六七八九十'):
                in_section = False
            if not in_section:
                continue
            if re.match(r'^[12]、', full) and not loc.startswith('table'):
                if item_idx < len(all_items):
                    _set_runs_text(p, f'{item_idx+1}、{all_items[item_idx]}；', True)
                    item_idx += 1
                    last_detection_p = p
                else:
                    _set_runs_text(p, '', True)
            elif '本工程设计文件及相关要求' in full or full.startswith('3、'):
                _set_runs_text(p, f'{fixed_item_num}、本工程设计文件及相关要求。', True)
                break

        # 额外项插入
        if len(all_items) > 2 and last_detection_p is not None:
            from docx.oxml import OxmlElement as _O
            ref = last_detection_p._element
            for extra_idx in range(2, len(all_items)):
                new_p = _O('w:p')
                ref.addnext(new_p)
                ref = new_p
                for pp in doc.paragraphs:
                    if pp._element == new_p:
                        run = pp.add_run(f'{extra_idx+1}、{all_items[extra_idx]}；')
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        break

    # ===== 11. 仪器表格（Table 3 — docx中可能是 Table 4 或动态匹配）=====
    instruments = data.get('instruments', [])
    if instruments:
        # 找仪器表：表头包含 "仪器名称" 或 "编号" 或 "检定"
        inst_table = None
        for ti, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                first_row_texts = ''.join([cell.text.strip() for cell in table.rows[0].cells])
                if '仪器名称' in first_row_texts and '编号' in first_row_texts:
                    inst_table = table
                    break

        if inst_table is not None:
            header_rows = 1
            needed = len(instruments) + header_rows
            while len(inst_table.rows) < needed:
                inst_table.add_row()
            while len(inst_table.rows) > needed:
                tr = inst_table.rows[-1]._tr
                inst_table._tbl.remove(tr)
            for ri in range(len(instruments)):
                row = inst_table.rows[ri + 1]
                inst = instruments[ri]
                for ci, col_key in enumerate(['name', 'number', 'calib_date', 'cert_number']):
                    if ci < len(row.cells):
                        for p in row.cells[ci].paragraphs:
                            _set_runs_text(p, str(inst.get(col_key, '')), True)

    # ===== 12. 表8 — 原始数据表（动态匹配表头）=====
    raw_data = data.get('raw_data', [])
    sample_count = int(data.get('sample_count', 0) or 0)
    if sample_count > 0:
        while len(raw_data) < sample_count:
            raw_data.append({'point_id': '', 'depth': '', 'blows': ''})

    if raw_data:
        # 找表8：表头包含"点号"但不含"承载力"（排除表9）
        raw_table = None
        for ti, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                hr = ''.join([cell.text.strip().replace(' ', '') for cell in table.rows[0].cells])
                if '点号' in hr and '承载力' not in hr and '土层' not in hr:
                    raw_table = table
                    break

        if raw_table is not None:
            while len(raw_table.rows) > 1:
                tr = raw_table.rows[-1]._tr
                raw_table._tbl.remove(tr)
            for rd in raw_data:
                new_row = raw_table.add_row()
                for ci, col_key in enumerate(['point_id', 'depth', 'blows']):
                    if ci < len(new_row.cells):
                        _safe_set_cell_text(new_row.cells[ci], str(rd.get(col_key, '')))
            # 设置行高1cm (1cm = 28.35磅)
            for row in raw_table.rows:
                row.height = Pt(28.35)
                for ci, cell in enumerate(row.cells):
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = etree2.SubElement(tc, qn('w:tcPr'))
                    vAlign = tcPr.find(qn('w:vAlign'))
                    if vAlign is None:
                        vAlign = etree2.SubElement(tcPr, qn('w:vAlign'))
                    vAlign.set(qn('w:val'), 'center')
                    for p in cell.paragraphs:
                        if ci == 2:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT   # 第三列(锤击数)左对齐
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 其他列居中

    # ===== 13. 表9 — 汇总表 =====
    summary_data = data.get('summary_data', [])
    sample_count = int(data.get('sample_count', 0) or 0)
    if sample_count > 0:
        while len(summary_data) < sample_count:
            summary_data.append({'soil_layer': '素填土', 'point_id': '', 'elevation': '', 'avg_blows': '', 'bearing_capacity': ''})

    if summary_data:
        sum_table = None
        for ti, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                hr = ''.join([cell.text.strip().replace(' ', '') for cell in table.rows[0].cells])
                if '土层' in hr and '承载力' in hr:
                    sum_table = table
                    break

        if sum_table is not None:
            while len(sum_table.rows) > 1:
                tr = sum_table.rows[-1]._tr
                sum_table._tbl.remove(tr)

            for sd in summary_data:
                new_row = sum_table.add_row()
                soil_name = str(sd.get('soil_layer', '素填土') or '素填土')
                c0 = new_row.cells[0]
                for p in c0.paragraphs:
                    for r in p.runs:
                        _safe_set_run_text(r, '')
                run_soil = c0.paragraphs[0].add_run(soil_name)
                run_soil.font.size = Pt(10.5)
                c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if len(new_row.cells) > 1:
                    _safe_set_cell_text(new_row.cells[1], str(sd.get('point_id', '')))
                if len(new_row.cells) > 2:
                    _safe_set_cell_text(new_row.cells[2], str(sd.get('elevation', '')))
                if len(new_row.cells) > 3:
                    _safe_set_cell_text(new_row.cells[3], str(sd.get('avg_blows', '')))
                if len(new_row.cells) > 4:
                    _safe_set_cell_text(new_row.cells[4], str(sd.get('bearing_capacity', '')))

            # 设置列宽
            tbl_grid = sum_table._tbl.find(qn('w:tblGrid'))
            if tbl_grid is not None:
                grid_cols = tbl_grid.findall(qn('w:gridCol'))
                col_widths = [1134, 1418, 1418, 1418, 1418]
                for i, w in enumerate(col_widths):
                    if i < len(grid_cols):
                        grid_cols[i].set(qn('w:w'), str(w))

            # 设置行高1cm (1cm = 28.35磅)
            for row in sum_table.rows:
                row.height = Pt(28.35)
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = etree2.SubElement(tc, qn('w:tcPr'))
                    noWrap = tcPr.find(qn('w:noWrap'))
                    if noWrap is None:
                        etree2.SubElement(tcPr, qn('w:noWrap'))
                    vAlign = tcPr.find(qn('w:vAlign'))
                    if vAlign is None:
                        vAlign = etree2.SubElement(tcPr, qn('w:vAlign'))
                    vAlign.set(qn('w:val'), 'center')
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 第一列（土层）合并：连续相同土层合并为一个单元格
            soil_values = [str(sd.get('soil_layer', '素填土') or '素填土') for sd in summary_data]
            if len(soil_values) > 1:
                group_start = 0
                for i in range(1, len(soil_values)):
                    if soil_values[i] != soil_values[group_start]:
                        # 结束当前分组
                        if i - group_start > 1:
                            _merge_table9_soil_cells(sum_table, group_start, i - 1)
                        group_start = i
                # 处理最后一个分组
                if len(soil_values) - group_start > 1:
                    _merge_table9_soil_cells(sum_table, group_start, len(soil_values) - 1)

            # ===== 13b. 表9分页控制 — 标题+表格一起固定在第十二页 =====
            # 找到表9前的"地基土承载力确定表"段落，在它之前插入分页符，
            # 让"地基土承载力确定表"+"表9"+表格一起从第12页开始。
            body = doc.element.body
            body_children = list(body)

            tbl_elem = sum_table._tbl
            tbl_body_idx = body_children.index(tbl_elem)

            # 往前找"地基土承载力确定表"（完整的表格标题段落）
            title_idx = None
            for k in range(tbl_body_idx - 1, -1, -1):
                child = body_children[k]
                if not child.tag.endswith('}p'):
                    continue
                txt = etree2.tostring(child, encoding='unicode')
                clean = re.sub(r'<[^>]+>', '', txt).strip()
                # 精确匹配表格标题："地基土承载力确定表" 或 "表  9"（纯编号）
                if '地基土承载力确定表' in clean:
                    title_idx = k
                    break
                # "表  9" / "表9"（纯编号，不含其他文字如","）
                if re.match(r'^表\s*9$', clean):
                    title_idx = k
                    # 继续往前找"地基土承载力确定表"
                    continue
                if clean and '表' not in clean and '地基土' not in clean:
                    break

            # 创建分页符段落
            pb = doc.add_paragraph()
            pb.add_run('')
            br = etree2.SubElement(pb.runs[0]._element, qn('w:br'))
            br.set(qn('w:type'), 'page')

            if title_idx is not None:
                body.remove(pb._element)
                body.insert(title_idx, pb._element)
            else:
                body.remove(pb._element)
                body.insert(tbl_body_idx, pb._element)

            # 不在表9后强制分页，让"九、结论、建议"可紧跟表9后面
            # "十、附图"的分页由 13c 控制

            # ===== 13c. "十、附图"分页控制 — 固定在第十三页开头 =====
            # 在"十、附图"段落前插入分页符，让该段落及其下所有内容（含附图说明）从第13页开始
            body = doc.element.body
            body_children = list(body)

            fu_idx = None
            for k, child in enumerate(body_children):
                if child.tag.endswith('}p'):
                    txt = etree2.tostring(child, encoding='unicode')
                    clean = re.sub(r'<[^>]+>', '', txt).strip()
                    if clean == '十、附图':
                        fu_idx = k
                        break

            if fu_idx is not None:
                pb3 = doc.add_paragraph()
                pb3.add_run('')
                br3 = etree2.SubElement(pb3.runs[0]._element, qn('w:br'))
                br3.set(qn('w:type'), 'page')
                body.remove(pb3._element)
                body.insert(fu_idx, pb3._element)

    # ===== 14. 检测结论（首页正文 + 第九章）=====
    conclusion = data.get('test_conclusion', '')
    # 第九章正文中的结论段落（包含"动力触探试验结果统计显示"）
    for loc, p in all_p:
        full = p.text.strip()
        if '动力触探试验结果统计显示' in full and '换算得出' in full and loc.startswith('body'):
            reds = _red_runs(p)
            if reds:
                for i, (ri, run, _) in enumerate(reds):
                    _safe_set_run_text(run, conclusion if i == 0 else '')
                    _safe_set_font_color(run, BLACK)

    # ===== 15. 结论、建议（第九章）=====
    suggestion_on = data.get('suggestion_on', True)
    suggestion_type = data.get('suggestion_type', 'qualified')

    if suggestion_type == 'qualified':
        suggestion_content = "2、基础施工过程中，望有关部门加强截排水及验槽工作。"
    else:
        suggestion_content = "2、建议对不满足设计要求的地基采取有效方式进行相应处理后再进行下一步施工。"

    # 15a. 更新目录页（第5页）TOC 条目
    # TOC 条目是 PAGEREF 字段，文本在 <w:hyperlink> 内的 <w:t> 中，
    # python-docx 的 paragraph.runs 无法直接访问。
    # 当不勾选"包含建议章节"时，把 "九、结论、建议" 改为 "九、结论"，
    # 并同步补长引导点（……），让右边页码对齐。
    if not suggestion_on:
        for p in doc.paragraphs:
            full = p.text.strip()
            if '九' in full and '结论' in full:
                is_toc = any(
                    'PAGEREF' in (instr.text or '')
                    for instr in p._element.iter(qn('w:instrText'))
                )
                if is_toc:
                    # 找到标题文本的 <w:t> 和引导点的 <w:t>
                    t_elems = list(p._element.iter(qn('w:t')))
                    title_elem = None
                    dot_elem = None
                    for idx, t_elem in enumerate(t_elems):
                        if '九、结论' in (t_elem.text or ''):
                            title_elem = t_elem
                            # 引导点通常在标题 <w:t> 的下一个 <w:t>
                            if idx + 1 < len(t_elems):
                                dot_elem = t_elems[idx + 1]
                            break
                    if title_elem is not None:
                        old_title = title_elem.text  # e.g. "九、结论、建议"
                        new_title = '九、结论'
                        title_elem.text = new_title
                        # 补长引导点：短了多少字符就补多少个 "…"
                        if old_title and dot_elem is not None and dot_elem.text:
                            delta = len(old_title) - len(new_title)
                            if delta > 0:
                                dot_elem.text = dot_elem.text + '…' * delta
                    break  # 只改第一个匹配的 TOC 条目

    # 辅助函数：判断段落是否为 TOC/目录段落
    def _is_toc_paragraph(para):
        """判断是否为目录（TOC/PAGEREF）段落"""
        # 方法1：段落样式包含 TOC
        try:
            if para.style and para.style.name and 'TOC' in para.style.name.upper():
                return True
        except Exception:
            pass
        # 方法2：XML 中包含 PAGEREF 字段指令
        return any(
            'PAGEREF' in (instr.text or '')
            for instr in para._element.iter(qn('w:instrText'))
        )

    section9_title_found = False
    section9_conclusion_filled = False

    for loc, p in all_p:
        full = p.text.strip()
        # 跳过目录/TOC 段落（PAGEREF 不在 p.text 中显现，需用 XML 检测）
        if _is_toc_paragraph(p):
            continue
        no_space = full.replace(' ', '').replace('\u3000', '')
        if no_space.startswith('九') and '结论' in no_space:
            section9_title_found = True
            for i, run in enumerate(p.runs):
                title_text = '九、结论、建议' if suggestion_on else '九、结论'
                _safe_set_run_text(run, title_text if i == 0 else '')
                _safe_set_font_color(run, BLACK)
            continue
        if section9_title_found and not section9_conclusion_filled and not full.startswith('2'):
            # 检测结论段落已由 section 14 处理（只填红色 run，保留黑色前缀）
            # 此处仅标记已完成，不再重复填充，避免破坏前缀文本
            section9_conclusion_filled = True
            continue
        if section9_title_found and full.startswith('2') and ('望有关部门' in full or '基础施工' in full or '建议对不满足' in full):
            if suggestion_on:
                reds = _red_runs(p)
                if reds:
                    for i, (ri, run, _) in enumerate(reds):
                        _safe_set_run_text(run, suggestion_content if i == 0 else '')
                        _safe_set_font_color(run, BLACK)
                else:
                    _set_runs_text(p, suggestion_content, False)
            else:
                for run in p.runs:
                    _safe_set_run_text(run, '')
                    _safe_set_font_color(run, BLACK)

    # ===== 16. 附图插入 =====
    images = data.get('images', [])
    if images:
        from docx.shared import Inches
        target_para = None
        for p in doc.paragraphs:
            if p.text and '十' in p.text and '附图' in p.text:
                target_para = p
                break
        if target_para is None:
            target_para = doc.add_paragraph()
            _safe_set_paragraph_text(target_para, '十、附图')

        for img_info in images:
            img_path = img_info.get('path', '')
            caption = img_info.get('caption', '')
            if img_path and os.path.exists(img_path):
                if caption:
                    cap_p = doc.add_paragraph()
                    cap_run = cap_p.add_run(caption)
                    cap_run.bold = True
                    cap_run.font.size = Pt(10.5)
                img_p = doc.add_paragraph()
                img_run = img_p.add_run()
                try:
                    img_run.add_picture(img_path, width=Inches(5.5))
                except Exception:
                    pass
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 17. 标题居中 =====
    for p in doc.paragraphs:
        if p.text and '圆锥动力触探试验' in p.text and '检测报告' not in p.text:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    # ===== 18. 全局红色→黑色 =====
    for p in doc.paragraphs:
        try:
            runs = p.runs
        except AttributeError:
            continue  # 跳过裸 XML 段落（如分页符段落），无 run 结构
        for run in runs:
            try:
                if run.font.color and run.font.color.rgb and str(run.font.color.rgb).upper() == 'FF0000':
                    _safe_set_font_color(run, BLACK)
            except OSError:
                pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            if run.font.color and run.font.color.rgb and str(run.font.color.rgb).upper() == 'FF0000':
                                _safe_set_font_color(run, BLACK)
                        except OSError:
                            pass
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr:
                for p in hdr.paragraphs:
                    try:
                        runs = p.runs
                    except AttributeError:
                        continue
                    for run in runs:
                        try:
                            if run.font.color and run.font.color.rgb and str(run.font.color.rgb).upper() == 'FF0000':
                                _safe_set_font_color(run, BLACK)
                        except OSError:
                            pass

    # ===== 19. 检测深度间隔符号 =====
    test_depth = data.get('test_depth_range', '')
    for loc, p in all_p:
        full = p.text.strip()
        if re.match(r'^\d+\.\d+[～~]\d+\.\d+$', full):
            _set_runs_text(p, test_depth, True)

    doc.save(output_path)
    return output_path


def _apply_superscript(paragraph, text):
    """将段落文本中的 ^{...} 标记转换为上标，_{...} 标记转换为下标
    新添加的 run 会继承段落第一个有效 run 的字体格式（字号、字体名称等）
    """
    import re as _re
    from docx.shared import RGBColor
    
    # 匹配上标 ^{...} 或下标 _{...}
    pattern = r'(\^\{([^}]+)\}|_\{([^}]+)\})'
    if not _re.search(pattern, text):
        # 没有上标/下标标记，直接填充
        _set_runs_text(paragraph, text, True)
        return text

    # 获取参考格式（从第一个有效 run 中获取字体属性）
    ref_run = None
    for r in paragraph.runs:
        if r.text.strip():
            ref_run = r
            break
    if ref_run is None and len(paragraph.runs) > 0:
        ref_run = paragraph.runs[0]
    
    # 提取参考格式
    ref_size = None
    ref_name = None
    ref_bold = None
    ref_italic = None
    ref_color = None
    if ref_run:
        try:
            ref_size = ref_run.font.size
            ref_name = ref_run.font.name
            ref_bold = ref_run.font.bold
            ref_italic = ref_run.font.italic
            if ref_run.font.color and ref_run.font.color.rgb:
                ref_color = ref_run.font.color.rgb
        except:
            pass
    
    # 清空所有 run 的文本
    for r in paragraph.runs:
        r.text = ''

    def _apply_ref_format(run):
        """给 run 应用参考格式"""
        if ref_size:
            run.font.size = ref_size
        if ref_name:
            run.font.name = ref_name
        if ref_bold is not None:
            run.font.bold = ref_bold
        if ref_italic is not None:
            run.font.italic = ref_italic
        if ref_color:
            run.font.color.rgb = ref_color
    
    idx = 0
    for m in _re.finditer(pattern, text):
        prefix = text[idx:m.start()]
        if prefix:
            paragraph.runs[0].text = (paragraph.runs[0].text or '') + prefix
            _apply_ref_format(paragraph.runs[0])
        
        # 判断是上标还是下标
        if m.group(1).startswith('^'):
            # 上标 ^{...}
            script_text = m.group(2)
            r_script = paragraph.add_run(script_text)
            r_script.font.superscript = True
        else:
            # 下标 _{...}
            script_text = m.group(3)
            r_script = paragraph.add_run(script_text)
            r_script.font.subscript = True
        
        _apply_ref_format(r_script)
        idx = m.end()

    tail = text[idx:]
    if tail:
        r_tail = paragraph.add_run(tail)
        _apply_ref_format(r_tail)

    # 返回去除标记后的纯文本（用于日志或调试）
    result = _re.sub(r'\^\{([^}]+)\}', r'\1', text)
    result = _re.sub(r'_\{([^}]+)\}', r'\1', result)
    return result


def _safe_log(msg):
    try:
        import sys
        sys.stderr.buffer.write((msg + '\n').encode('utf-8', 'replace'))
        sys.stderr.buffer.flush()
    except Exception:
        pass


def _kill_word():
    try:
        subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
                       capture_output=True, timeout=5)
        time.sleep(2)
    except Exception:
        pass
